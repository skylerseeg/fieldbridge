"""
Bid Intelligence Service
PDF/DOCX ingestion pipeline for drawings and spec books.
Extracts structured text → feeds to bid_agent for BOM extraction.
"""
import logging
import tempfile
from pathlib import Path
from typing import Union

log = logging.getLogger("fieldbridge.bid_intelligence")


def extract_text_from_pdf(source: Union[str, bytes]) -> str:
    """Extract all text from a PDF, preserving table structure."""
    import pdfplumber

    if isinstance(source, bytes):
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(source)
            source = tmp.name

    pages = []
    with pdfplumber.open(source) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            for table in (page.extract_tables() or []):
                for row in table:
                    if row:
                        text += "\n" + " | ".join(str(c) for c in row if c)
            pages.append(f"--- PAGE {i+1} ---\n{text}")

    return "\n\n".join(pages)


def extract_text_from_docx(source: Union[str, bytes]) -> str:
    """Extract text from a Word document."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(source) if isinstance(source, bytes) else source)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines)


def chunk_document(text: str, max_chars: int = 12000) -> list[str]:
    """Split large document text into Claude-friendly chunks at page boundaries."""
    pages = text.split("--- PAGE ")
    chunks: list[str] = []
    current = ""

    for page in pages:
        segment = ("--- PAGE " if page and not page.startswith("---") else "") + page
        if len(current) + len(segment) < max_chars:
            current += segment
        else:
            if current:
                chunks.append(current.strip())
            current = segment

    if current:
        chunks.append(current.strip())

    return chunks or [text[:max_chars]]
